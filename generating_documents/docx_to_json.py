import json
from docx import Document

def docx_to_json(docx_file):
    doc = Document(docx_file)

    data = {}

    for i in range(len(doc.paragraphs)):
        paragraph = doc.paragraphs[i]

        if paragraph.style.name == 'Heading 1':
            heading = paragraph.text
            next_paragraph = doc.paragraphs[i+1].text
            if heading == 'Age':
                next_paragraph = int(next_paragraph)
            data[heading] = next_paragraph

    return data



if __name__ == '__main__':
    docx_path = "resources/test_person.docx"

    json_data = docx_to_json(docx_path)

    json_path = "resources/file.json"
    with open(json_path, 'w') as json_file:
        json.dump(json_data, json_file, indent=4)