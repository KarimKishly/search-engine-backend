import json
from docx import Document

def json_to_docx(element, index):
    doc = Document()

    doc.add_heading('Name', level=1)
    doc.add_paragraph(element['Name'])

    doc.add_heading('Gender', level=1)
    doc.add_paragraph(element['Gender'])

    doc.add_heading('Age', level=1)
    doc.add_paragraph(str(element['Age']))

    doc.add_heading('Date of Birth', level=1)
    doc.add_paragraph(element['Date of Birth'])

    doc.add_heading('Description', level=1)
    doc.add_paragraph(element['Description'])

    doc.add_heading('Occupation', level=1)
    doc.add_paragraph(element['Occupation'])

    doc.add_heading('Interests', level=1)
    doc.add_paragraph(element['Interests'])

    doc.add_heading('Likes', level=1)
    doc.add_paragraph(str(element['Likes']))

    doc.save(f'word_documents/person_{index}.docx')

if __name__ == '__main__':
    with open('resources/people_dataset.json', 'r') as file:
        anime_json = json.load(file)
    index = 1
    for element in anime_json:
        json_to_docx(element, index)
        index+=1